from typing import Any, Optional
from docx.opc.exceptions import PackageNotFoundError
from pdfminer.psparser import PSException
from docx.api import Document
from zipfile import ZipFile
from app.exceptions import NoPropertiesError
from app.models.content import Content
from app.models.handlers import FileHandler


class DOCXHandler(FileHandler):

    def extract_metadata(self) -> Optional[dict[str, Any]]:
        # TODO: figure out why the metadata came empty in the test files.
        try:
            prop = Document(str(self.file)).core_properties
            metadata = {}
            for d in dir(prop):
                if not d.startswith('_'):
                    metadata[d] = getattr(prop, d)
            return metadata
        except PackageNotFoundError:
            raise FileNotFoundError
        except PSException as e:
            self.logger.error(e)

    def extract_content(self) -> Content:
        def _extract_text(doc) -> str:
            return '\n'.join([para.text for para in doc.paragraphs])

        def _extract_tables(doc) -> list[list[list[str]]]:
            tables = []
            for table in doc.tables:
                single_table = []
                for row in table.rows:
                    # for every cell of the table, get the text on the paragraph and make it part of a single string
                    content = [", ".join([p.text for p in i.paragraphs])
                               for i in row.cells]
                    single_table.append(content)
                tables.append(single_table)
            return tables

        doc = Document(str(self.file))
        text = _extract_text(doc)
        tables = _extract_tables(doc)
        return Content(text, tables)

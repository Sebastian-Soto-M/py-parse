from collections import defaultdict
import zipfile
import docx
from typing import Dict
import xml.etree.ElementTree as ET
from typing import Optional, List, Dict, Any
import logging
from abc import ABC, abstractmethod
from pathlib import Path
from typing import Any, Optional

import pdfplumber
from app.exceptions import NoPropertiesError

from app.models.content import Content, ContentWithImage


class FileHandler(ABC):
    def __init__(self, filename: str):
        self.file = Path(filename)
        self.logger = logging.getLogger(self.__class__.__name__)

    @abstractmethod
    def extract_metadata(self) -> Optional[dict[str, Any]]:
        raise NotImplementedError

    @abstractmethod
    def extract_content(self) -> Content:
        raise NotImplementedError


class ImageHandler(FileHandler):
    def extract_metadata(self):
        # Logic to extract image metadata
        return "Image Metadata"

    def extract_content(self):
        # Logic to extract image content
        return "Image Content"


class DOCXHandler(FileHandler):

    def extract_metadata(self) -> Optional[dict[str, Any]]:
        try:
            metadata = {}
            with zipfile.ZipFile(self.file, 'r') as docx:
                with docx.open('docProps/core.xml') as file:
                    tree = ET.parse(file)
                    root = tree.getroot()

                    for element in root:
                        for key in ns.keys():
                            ns_key = f"{{{ns[key]}}}"
                            if ns_key in element.tag:
                                tag = element.tag.split(ns_key)[1]
                                metadata[tag] = element.text
        except KeyError as e:
            raise NoPropertiesError(e)
        return metadata

    def extract_content(self) -> Content:
        def _extract_text(doc) -> str:
            return '\n'.join([para.text for para in doc.paragraphs])

        def _extract_tables(doc) -> list:
            tables = []
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    # this character will be the column divider.
                    row_content = "|".join([cell.text for cell in row.cells])
                    table_content.append(row_content)
                tables.append(table_content)
            return tables

        doc = docx.Document(str(self.file))
        text = _extract_text(doc)
        tables = _extract_tables(doc)
        return Content(text, tables)


class PDFHandler(FileHandler):

    def _open_pdf(self) -> Optional[pdfplumber.PDF]:
        """
        Safely opens a PDF file.
        Returns a pdfplumber.PDF object or None if an error occurs.
        """
        try:
            return pdfplumber.open(self.file)
        except (FileNotFoundError, IOError) as e:
            logging.error(f"Error opening PDF file: {e}")
            return None

    def _parse_date(self, date):
        # Original date example = "D:20230116165046+01'00'"
        # Format the date string by removing 'D:' and replacing "'" with ""
        formatted_date = date[2:][:8]
        # Parse and convert to ISO 8601 format
        try:
            return datetime.strptime(formatted_date, '%Y%m%d').isoformat()
        except Exception as e:
            self.logger.error(e)
            return ""

    def extract_metadata(self) -> dict[str, Any]:
        """
        Extracts metadata from the PDF file.
        Returns a dictionary containing metadata.
        """
        try:
            with self._open_pdf() as pdf:
                if isinstance(pdf, pdfplumber.PDF):
                    return {key: self._parse_date(value) if "date" in key.lower() else value
                            for key, value in pdf.metadata.items()}
        except Exception as e:
            self.logger.error(e)
        return {}

    def extract_content(self) -> Content:
        """
        Extracts text, table, and image data from the PDF file.
        Returns a PDFContent object.
        """
        with self._open_pdf() as pdf:
            if pdf:
                text_content = self._extract_text_content(pdf)
                table_data = self._extract_table_data(pdf)
                image_data = self._extract_image_data(pdf)

                return ContentWithImage(text_content, table_data, image_data)
            return Content("", [])

    def _extract_text_content(self, pdf: pdfplumber.PDF) -> str:
        """ Extracts and returns textual content from the PDF. """
        return "\n".join(page.extract_text() or "" for page in pdf.pages).strip()

    def _extract_table_data(self, pdf: pdfplumber.PDF) -> List[List[List[str]]]:
        """ Extracts and returns table data from the PDF. """
        all_tables = []
        for page in pdf.pages:
            tables = page.extract_tables()
            if tables:
                all_tables.extend(tables)
        return all_tables

    def _extract_image_data(self, pdf: pdfplumber.PDF) -> List[Dict[str, Any]]:
        """ Extracts and returns image data from the PDF. """
        return [img for page in pdf.pages for img in page.images]
